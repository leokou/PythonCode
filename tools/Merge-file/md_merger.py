# -*- coding: utf-8 -*-
"""
文件合并工具 - 将多个 txt / md / docx 文件拖拽合并为一个 Markdown 文件
"""

import os
import sys
import threading
import tkinter as tk
from tkinter import ttk, filedialog, messagebox

# 尝试导入拖拽库
try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
    DND_AVAILABLE = True
except ImportError:
    DND_AVAILABLE = False

# 尝试导入 docx 库
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# 尝试导入 PyGithub 库
try:
    from github import Github
    GITHUB_AVAILABLE = True
except ImportError:
    GITHUB_AVAILABLE = False


SUPPORTED_EXT = {'.txt', '.md', '.markdown', '.docx'}


def read_txt_file(filepath):
    """读取 txt / md 文件内容"""
    encodings = ['utf-8', 'utf-8-sig', 'gbk', 'gb2312', 'latin-1']
    for enc in encodings:
        try:
            with open(filepath, 'r', encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    # 全部失败时用二进制读再容错
    with open(filepath, 'rb') as f:
        return f.decode('utf-8', errors='replace')


def read_docx_file(filepath):
    """读取 docx 文件内容，转为 Markdown 格式"""
    if not DOCX_AVAILABLE:
        return "[警告：未安装 python-docx，无法读取 Word 文件，请先 pip install python-docx]"

    doc = Document(filepath)
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append('')
            continue

        # 处理标题样式
        style_name = para.style.name if para.style else ''
        if style_name.startswith('Heading 1') or style_name == '标题 1':
            lines.append(f'# {text}')
        elif style_name.startswith('Heading 2') or style_name == '标题 2':
            lines.append(f'## {text}')
        elif style_name.startswith('Heading 3') or style_name == '标题 3':
            lines.append(f'### {text}')
        elif style_name.startswith('Heading') or style_name.startswith('标题'):
            lines.append(f'#### {text}')
        else:
            # 普通段落
            lines.append(text)

    # 处理表格
    for table in doc.tables:
        lines.append('')
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip().replace('|', '\\|') for cell in row.cells]
            lines.append('| ' + ' | '.join(cells) + ' |')
            if i == 0:
                lines.append('| ' + ' | '.join(['---'] * len(cells)) + ' |')
        lines.append('')

    return '\n'.join(lines)


def read_file_content(filepath):
    """根据扩展名读取文件内容"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.docx':
        return read_docx_file(filepath)
    else:
        return read_txt_file(filepath)


class FileMergerApp:
    def __init__(self, root):
        self.root = root
        self.root.title('文件合并工具 - Markdown')
        self.root.geometry('720x560')
        self.root.minsize(600, 480)

        self.file_list = []  # 存储文件路径列表
        self.last_output_path = None  # 最后一次合并的输出文件路径

        self._build_ui()
        self._bind_dnd()

        # 底部状态栏
        self.status_var = tk.StringVar(value=f'就绪 | 支持格式: .txt .md .docx | 拖拽库: {"已启用" if DND_AVAILABLE else "未安装"}')
        status_bar = ttk.Label(root, textvariable=self.status_var, relief='sunken', anchor='w')
        status_bar.pack(side='bottom', fill='x')

    def _build_ui(self):
        # 顶部说明区
        top_frame = ttk.Frame(self.root, padding=10)
        top_frame.pack(fill='x')

        title = ttk.Label(top_frame, text='📄 文件合并工具', font=('Microsoft YaHei', 14, 'bold'))
        title.pack(side='left')

        hint = ttk.Label(top_frame, text='将 txt / md / docx 文件拖拽到下方列表，按顺序合并为一个 Markdown 文件',
                         foreground='#666')
        hint.pack(side='left', padx=15)

        # 主区域：文件列表
        main_frame = ttk.Frame(self.root, padding=(10, 0, 10, 10))
        main_frame.pack(fill='both', expand=True)

        # Listbox + Scrollbar
        list_frame = ttk.Frame(main_frame)
        list_frame.pack(fill='both', expand=True, side='left')

        self.listbox = tk.Listbox(list_frame, selectmode=tk.EXTENDED,
                                  font=('Consolas', 10), activestyle='dotbox')
        self.listbox.pack(side='left', fill='both', expand=True)

        scrollbar = ttk.Scrollbar(list_frame, orient='vertical', command=self.listbox.yview)
        scrollbar.pack(side='right', fill='y')
        self.listbox.config(yscrollcommand=scrollbar.set)

        # 右侧按钮区
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(side='right', fill='y', padx=(10, 0))

        ttk.Button(btn_frame, text='➕ 添加文件', command=self.add_files).pack(fill='x', pady=2)
        ttk.Button(btn_frame, text='📁 添加文件夹', command=self.add_folder).pack(fill='x', pady=2)
        ttk.Button(btn_frame, text='🗑 移除选中', command=self.remove_selected).pack(fill='x', pady=2)
        ttk.Button(btn_frame, text='🧹 清空列表', command=self.clear_all).pack(fill='x', pady=2)

        ttk.Separator(btn_frame, orient='horizontal').pack(fill='x', pady=8)

        ttk.Button(btn_frame, text='⬆ 上移', command=self.move_up).pack(fill='x', pady=2)
        ttk.Button(btn_frame, text='⬇ 下移', command=self.move_down).pack(fill='x', pady=2)

        ttk.Separator(btn_frame, orient='horizontal').pack(fill='x', pady=8)

        merge_btn = ttk.Button(btn_frame, text='🚀 合并导出', command=self.do_merge)
        merge_btn.pack(fill='x', pady=2)
        merge_btn.state(['!disabled'])

        ttk.Separator(btn_frame, orient='horizontal').pack(fill='x', pady=8)

        github_btn = ttk.Button(btn_frame, text='☁️ 上传到 GitHub', command=self.do_upload_github)
        github_btn.pack(fill='x', pady=2)
        merge_github_btn = ttk.Button(btn_frame, text='🔄 合并同步GitHub', command=self.do_merge_and_upload)
        merge_github_btn.pack(fill='x', pady=2)

        # 拖拽提示覆盖层（仅在无文件时显示提示文字）
        self.drop_hint = ttk.Label(self.listbox,
                                    text='\n\n  将文件拖拽到这里\n\n  或点击「添加文件」按钮  \n\n  支持 .txt / .md / .docx ',
                                    foreground='#999', font=('Microsoft YaHei', 11),
                                    justify='center')
        self._show_hint(True)

    def _show_hint(self, show):
        if show:
            self.drop_hint.place(relx=0.5, rely=0.5, anchor='center')
        else:
            self.drop_hint.place_forget()

    def _bind_dnd(self):
        """绑定拖拽事件"""
        if DND_AVAILABLE:
            self.listbox.drop_target_register(DND_FILES)
            self.listbox.dnd_bind('<<Drop>>', self._on_drop)
            # 整个窗口也支持拖
            self.root.drop_target_register(DND_FILES)
            self.root.dnd_bind('<<Drop>>', self._on_drop)

    def _on_drop(self, event):
        """处理拖拽放下事件"""
        # tkinterdnd2 返回的 data 可能用 {} 包裹带空格的路径
        raw = event.data
        files = self._parse_drop_files(raw)
        self._add_files(files)

    def _parse_drop_files(self, raw):
        """解析拖拽的文件路径字符串"""
        files = []
        buf = ''
        in_brace = False
        for ch in raw:
            if ch == '{':
                in_brace = True
                buf = ''
            elif ch == '}':
                in_brace = False
                if buf:
                    files.append(buf)
                    buf = ''
            elif ch == ' ' and not in_brace:
                if buf:
                    files.append(buf)
                    buf = ''
            else:
                buf += ch
        if buf:
            files.append(buf)
        return files

    def add_files(self):
        """通过对话框添加文件"""
        filetypes = [
            ('支持的文件', '*.txt *.md *.markdown *.docx'),
            ('文本文件', '*.txt'),
            ('Markdown 文件', '*.md *.markdown'),
            ('Word 文档', '*.docx'),
            ('所有文件', '*.*'),
        ]
        paths = filedialog.askopenfilenames(title='选择要合并的文件', filetypes=filetypes)
        if paths:
            self._add_files(list(paths))

    def add_folder(self):
        """通过文件夹选择对话框添加文件"""
        folder_path = filedialog.askdirectory(title='选择文件夹')
        if not folder_path:
            return

        files = self._scan_folder(folder_path)
        if not files:
            messagebox.showinfo('提示', '该文件夹中没有支持的文件类型')
            return

        files_sorted = sorted(files, key=lambda x: (os.path.splitext(x)[1].lower(), x))

        self._show_folder_file_picker(files_sorted, folder_path)

    def _scan_folder(self, folder_path):
        """递归扫描文件夹，返回所有支持的文件"""
        result = []
        for root_dir, _, files in os.walk(folder_path):
            for fn in files:
                ext = os.path.splitext(fn)[1].lower()
                if ext in SUPPORTED_EXT:
                    result.append(os.path.join(root_dir, fn))
        return result

    def _show_folder_file_picker(self, files, folder_path):
        """显示文件夹文件选择对话框"""
        dialog = tk.Toplevel(self.root)
        dialog.title(f'选择文件 - {os.path.basename(folder_path)}')
        dialog.geometry('800x550')
        dialog.resizable(True, True)
        dialog.transient(self.root)
        dialog.grab_set()

        top_frame = ttk.Frame(dialog, padding=10)
        top_frame.pack(fill='x')

        ttk.Label(top_frame, text=f'📂 文件夹: {folder_path}', font=('Microsoft YaHei', 11)).pack(side='left')
        ttk.Label(top_frame, text=f'共 {len(files)} 个文件', foreground='#666').pack(side='right')

        btn_frame = ttk.Frame(dialog, padding=(10, 0, 10, 5))
        btn_frame.pack(fill='x')

        ttk.Button(btn_frame, text='全选', command=lambda: self._toggle_all_checkboxes(True, checkboxes)).pack(side='left', padx=2)
        ttk.Button(btn_frame, text='反选', command=lambda: self._invert_checkboxes(checkboxes)).pack(side='left', padx=2)

        list_frame = ttk.Frame(dialog, padding=(10, 0, 10, 10))
        list_frame.pack(fill='both', expand=True)

        canvas = tk.Canvas(list_frame)
        scrollbar = ttk.Scrollbar(list_frame, orient='vertical', command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)

        scrollable_frame.bind('<Configure>', lambda e: canvas.configure(scrollregion=canvas.bbox('all')))
        canvas.create_window((0, 0), window=scrollable_frame, anchor='nw')
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')

        def _on_mousewheel(e):
            canvas.yview_scroll(int(-1 * (e.delta / 120)), 'units')

        canvas.bind('<MouseWheel>', _on_mousewheel)

        checkboxes = []
        ext_order = {'.md': 0, '.markdown': 1, '.txt': 2, '.docx': 3}

        current_ext = None
        for filepath in files:
            ext = os.path.splitext(filepath)[1].lower()
            if ext != current_ext:
                current_ext = ext
                ext_label = ttk.Label(scrollable_frame, text=f'--- {ext.upper()} 文件 ---', font=('Microsoft YaHei', 10, 'bold'), foreground='#3366cc')
                ext_label.pack(anchor='w', pady=(5, 2))

            cb_var = tk.BooleanVar(value=True)
            cb = ttk.Checkbutton(scrollable_frame, text=filepath, variable=cb_var)
            cb.pack(anchor='w', padx=5)
            checkboxes.append((cb_var, filepath))

        bottom_frame = ttk.Frame(dialog, padding=(10, 5, 10, 10))
        bottom_frame.pack(fill='x')

        def on_ok():
            selected = [fp for var, fp in checkboxes if var.get()]
            if selected:
                self._add_files(selected)
                self._set_status(f'从文件夹添加了 {len(selected)} 个文件')
            dialog.destroy()

        ttk.Button(bottom_frame, text='确定添加', command=on_ok).pack(side='right', padx=5)
        ttk.Button(bottom_frame, text='取消', command=dialog.destroy).pack(side='right', padx=5)

    def _toggle_all_checkboxes(self, select_all, checkboxes):
        """全选或取消全选所有复选框"""
        for var, _ in checkboxes:
            var.set(select_all)

    def _invert_checkboxes(self, checkboxes):
        """反选所有复选框（取反）"""
        for var, _ in checkboxes:
            var.set(not var.get())

    def _add_files(self, paths):
        """添加文件到列表（去重、过滤格式）"""
        added = 0
        skipped = 0
        for p in paths:
            p = os.path.abspath(p)
            if not os.path.isfile(p):
                # 跳过目录
                if os.path.isdir(p):
                    # 递归添加目录下的支持文件
                    for root_dir, _, files in os.walk(p):
                        for fn in files:
                            ext = os.path.splitext(fn)[1].lower()
                            if ext in SUPPORTED_EXT:
                                full = os.path.join(root_dir, fn)
                                if full not in self.file_list:
                                    self.file_list.append(full)
                                    self.listbox.insert('end', full)
                                    added += 1
                continue

            ext = os.path.splitext(p)[1].lower()
            if ext not in SUPPORTED_EXT:
                skipped += 1
                continue
            if p in self.file_list:
                continue
            self.file_list.append(p)
            self.listbox.insert('end', p)
            added += 1

        if added > 0:
            self._show_hint(False)

        msg = f'已添加 {added} 个文件'
        if skipped > 0:
            msg += f'，跳过 {skipped} 个不支持的文件'
        self._set_status(msg)

    def remove_selected(self):
        """移除选中的文件"""
        sel = list(self.listbox.curselection())
        if not sel:
            return
        # 从后往前删，避免索引错乱
        for idx in reversed(sel):
            self.listbox.delete(idx)
            del self.file_list[idx]

        if len(self.file_list) == 0:
            self._show_hint(True)
        self._set_status(f'已移除 {len(sel)} 个文件，剩余 {len(self.file_list)} 个')

    def clear_all(self):
        """清空列表"""
        if not self.file_list:
            return
        if not messagebox.askyesno('确认', '确定要清空所有文件吗？'):
            return
        self.listbox.delete(0, 'end')
        self.file_list.clear()
        self._show_hint(True)
        self._set_status('已清空')

    def move_up(self):
        """选中项上移"""
        sel = list(self.listbox.curselection())
        if not sel or 0 in sel:
            return
        for idx in sel:
            if idx <= 0:
                continue
            # 交换数据
            self.file_list[idx - 1], self.file_list[idx] = self.file_list[idx], self.file_list[idx - 1]
            # 交换 listbox 显示
            val = self.listbox.get(idx)
            self.listbox.delete(idx)
            self.listbox.insert(idx - 1, val)
        # 重新选中
        new_sel = [i - 1 for i in sel]
        self.listbox.selection_clear(0, 'end')
        for i in new_sel:
            self.listbox.selection_set(i)

    def move_down(self):
        """选中项下移"""
        sel = list(self.listbox.curselection())
        last = len(self.file_list) - 1
        if not sel or last in sel:
            return
        # 从后往前处理
        for idx in reversed(sel):
            if idx >= last:
                continue
            self.file_list[idx + 1], self.file_list[idx] = self.file_list[idx], self.file_list[idx + 1]
            val = self.listbox.get(idx)
            self.listbox.delete(idx)
            self.listbox.insert(idx + 1, val)
        new_sel = [i + 1 for i in sel]
        self.listbox.selection_clear(0, 'end')
        for i in new_sel:
            self.listbox.selection_set(i)

    def do_merge(self):
        """执行合并"""
        if not self.file_list:
            messagebox.showwarning('提示', '请先添加文件')
            return

        output_path = filedialog.asksaveasfilename(
            title='保存合并后的 Markdown 文件',
            defaultextension='.md',
            initialfile='merged_output.md',
            filetypes=[('Markdown 文件', '*.md'), ('文本文件', '*.txt'), ('所有文件', '*.*')]
        )
        if not output_path:
            return

        try:
            self._merge_to_file(output_path)
            self.last_output_path = output_path
            self._set_status(f'合并完成 → {output_path}')
            if messagebox.askyesno('完成', f'合并成功！共 {len(self.file_list)} 个文件已合并。\n\n是否打开输出文件所在文件夹？'):
                self._open_folder(output_path)
        except Exception as e:
            messagebox.showerror('错误', f'合并失败：{str(e)}')
            self._set_status(f'合并失败: {e}')

    def do_upload_github(self):
        """上传列表中的所有文件到 GitHub"""
        if not self.file_list:
            messagebox.showwarning('提示', '请先添加文件')
            return

        if not GITHUB_AVAILABLE:
            messagebox.showerror('错误', '未安装 PyGithub，请先执行：\npip install pygithub')
            return

        github_dialog = tk.Toplevel(self.root)
        github_dialog.title('上传文件到 GitHub')
        github_dialog.geometry('400x300')
        github_dialog.resizable(False, False)
        github_dialog.transient(self.root)
        github_dialog.grab_set()

        ttk.Label(github_dialog, text='📤 上传文件到 GitHub', font=('Microsoft YaHei', 12, 'bold')).pack(pady=10)

        ttk.Label(github_dialog, text='GitHub Token:').pack(anchor='w', padx=20)
        token_var = tk.StringVar(value='')
        token_entry = ttk.Entry(github_dialog, textvariable=token_var, show='*', width=40)
        token_entry.pack(padx=20, pady=5)

        ttk.Label(github_dialog, text='仓库名 (如 leokou/leoshow):').pack(anchor='w', padx=20)
        repo_var = tk.StringVar(value='leokou/leoshow')
        repo_entry = ttk.Entry(github_dialog, textvariable=repo_var, width=40)
        repo_entry.pack(padx=20, pady=5)

        ttk.Label(github_dialog, text='分支名:').pack(anchor='w', padx=20)
        branch_var = tk.StringVar(value='main')
        branch_entry = ttk.Entry(github_dialog, textvariable=branch_var, width=40)
        branch_entry.pack(padx=20, pady=5)

        result_var = tk.StringVar(value=f'将上传 {len(self.file_list)} 个文件')
        result_label = ttk.Label(github_dialog, textvariable=result_var, foreground='green', wraplength=360)
        result_label.pack(pady=10)

        def on_upload():
            token = token_var.get().strip()
            repo_name = repo_var.get().strip()
            branch = branch_var.get().strip() or 'main'

            if not token:
                result_var.set('❌ 请输入 GitHub Token')
                return
            if not repo_name:
                result_var.set('❌ 请输入仓库名')
                return

            result_var.set(f'⏳ 正在上传 {len(self.file_list)} 个文件...')
            github_dialog.update()

            def upload_thread():
                success_count = 0
                fail_count = 0
                last_url = ''
                for filepath in self.file_list:
                    try:
                        success, msg = self._upload_to_github(filepath, token, repo_name, branch)
                        if success:
                            success_count += 1
                            last_url = msg
                        else:
                            fail_count += 1
                    except Exception as e:
                        fail_count += 1

                self.root.after(0, lambda: handle_result(success_count, fail_count, last_url))

            def handle_result(success_count, fail_count, last_url):
                if success_count > 0:
                    msg = f'✅ 上传完成！\n成功: {success_count} 个\n失败: {fail_count} 个'
                    if last_url:
                        msg += f'\n\n最后一个文件:\n{last_url}'
                        copied = self._copy_to_clipboard(last_url)
                        if copied:
                            msg += '\n\n📋 已复制到剪贴板'
                    result_var.set(msg)
                else:
                    result_var.set(f'❌ 全部上传失败')

            threading.Thread(target=upload_thread, daemon=True).start()

        ttk.Button(github_dialog, text='上传', command=on_upload).pack(pady=5)
        ttk.Button(github_dialog, text='关闭', command=github_dialog.destroy).pack(pady=5)

    def do_merge_and_upload(self):
        """合并文件并直接上传到 GitHub"""
        if not self.file_list:
            messagebox.showwarning('提示', '请先添加文件')
            return

        if not GITHUB_AVAILABLE:
            messagebox.showerror('错误', '未安装 PyGithub，请先执行：\npip install pygithub')
            return

        github_dialog = tk.Toplevel(self.root)
        github_dialog.title('合并同步到 GitHub')
        github_dialog.geometry('500x450')
        github_dialog.resizable(False, False)
        github_dialog.transient(self.root)
        github_dialog.grab_set()

        ttk.Label(github_dialog, text='🔄 合并同步到 GitHub', font=('Microsoft YaHei', 12, 'bold')).pack(pady=10)

        ttk.Label(github_dialog, text='输出文件名:').pack(anchor='w', padx=20)
        filename_var = tk.StringVar(value='merged_output.md')
        filename_entry = ttk.Entry(github_dialog, textvariable=filename_var, width=40)
        filename_entry.pack(padx=20, pady=5)

        ttk.Label(github_dialog, text='GitHub Token:').pack(anchor='w', padx=20)
        token_var = tk.StringVar(value='')
        token_entry = ttk.Entry(github_dialog, textvariable=token_var, show='*', width=40)
        token_entry.pack(padx=20, pady=5)

        ttk.Label(github_dialog, text='仓库名 (如 leokou/leoshow):').pack(anchor='w', padx=20)
        repo_var = tk.StringVar(value='leokou/leoshow')
        repo_entry = ttk.Entry(github_dialog, textvariable=repo_var, width=40)
        repo_entry.pack(padx=20, pady=5)

        ttk.Label(github_dialog, text='分支名:').pack(anchor='w', padx=20)
        branch_var = tk.StringVar(value='main')
        branch_entry = ttk.Entry(github_dialog, textvariable=branch_var, width=40)
        branch_entry.pack(padx=20, pady=5)

        result_var = tk.StringVar(value=f'将合并 {len(self.file_list)} 个文件并上传')
        result_label = ttk.Label(github_dialog, textvariable=result_var, foreground='green', wraplength=360)
        result_label.pack(pady=10)

        raw_url_var = tk.StringVar(value='')
        raw_url_label = ttk.Label(github_dialog, textvariable=raw_url_var, font=('Consolas', 9), wraplength=450, foreground='blue')
        raw_url_label.pack(pady=5)

        def on_merge_upload():
            filename = filename_var.get().strip()
            token = token_var.get().strip()
            repo_name = repo_var.get().strip()
            branch = branch_var.get().strip() or 'main'

            if not filename:
                result_var.set('❌ 请输入输出文件名')
                return
            if not token:
                result_var.set('❌ 请输入 GitHub Token')
                return
            if not repo_name:
                result_var.set('❌ 请输入仓库名')
                return

            result_var.set('⏳ 正在合并并上传...')
            github_dialog.update()

            def merge_upload_thread():
                try:
                    parts = []
                    for idx, filepath in enumerate(self.file_list, 1):
                        fname = os.path.basename(filepath)
                        content = read_file_content(filepath)
                        parts.append(f'## {idx}. {fname}\n\n')
                        parts.append(content)
                        parts.append('\n\n---\n\n')

                    merged_content = ''.join(parts)

                    g = Github(token)
                    repo = g.get_repo(repo_name)

                    try:
                        contents = repo.get_contents(filename, ref=branch)
                        repo.update_file(contents.path, f'Update {filename}', merged_content, contents.sha, branch=branch)
                    except Exception:
                        repo.create_file(filename, f'Create {filename}', merged_content, branch=branch)

                    raw_url = f'https://raw.githubusercontent.com/{repo_name}/refs/heads/{branch}/{filename}'
                    self.root.after(0, lambda: handle_result(True, raw_url))
                except Exception as e:
                    self.root.after(0, lambda: handle_result(False, str(e)))

            def handle_result(success, msg):
                if success:
                    result_var.set('✅ 合并同步成功！')
                    raw_url_var.set(msg)
                    merge_btn.pack_forget()
                    copy_btn.config(state='normal', bg='#4CAF50')
                else:
                    result_var.set(f'❌ 失败: {msg}')
                    raw_url_var.set('')
                    copy_btn.config(state='disabled', bg='#808080')

            threading.Thread(target=merge_upload_thread, daemon=True).start()

        def copy_raw():
            raw_url = raw_url_var.get()
            if raw_url.startswith('https://raw.'):
                self._copy_to_clipboard(raw_url)
                github_dialog.destroy()

        btn_frame = ttk.Frame(github_dialog)
        btn_frame.pack(pady=15)

        btn_style = {'font': ('Microsoft YaHei', 10), 'width': 14, 'height': 1}

        merge_btn = tk.Button(btn_frame, text='合并上传', command=on_merge_upload,
                              bg='#2196F3', fg='white', **btn_style)
        merge_btn.pack(side='left', padx=8)

        copy_btn = tk.Button(btn_frame, text='📋 复制Raw链接', command=copy_raw,
                             bg='#808080', fg='white', state='disabled', **btn_style)
        copy_btn.pack(side='left', padx=8)

        tk.Button(btn_frame, text='关闭', command=github_dialog.destroy,
                  bg='#607D8B', fg='white', **btn_style).pack(side='left', padx=8)

    def _merge_to_file(self, output_path):
        """核心合并逻辑"""
        parts = []
        for idx, filepath in enumerate(self.file_list, 1):
            filename = os.path.basename(filepath)
            content = read_file_content(filepath)
            # 每个文件前加标题分隔
            parts.append(f'## {idx}. {filename}\n\n')
            parts.append(content)
            parts.append('\n\n---\n\n')

        # 写文件
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(''.join(parts))

    def _open_folder(self, filepath):
        """打开文件所在文件夹"""
        folder = os.path.dirname(os.path.abspath(filepath))
        try:
            if sys.platform.startswith('win'):
                os.startfile(folder)
            elif sys.platform == 'darwin':
                os.system(f'open "{folder}"')
            else:
                os.system(f'xdg-open "{folder}"')
        except Exception:
            pass

    def _set_status(self, msg):
        self.status_var.set(msg)

    def _upload_to_github(self, file_path, token, repo_name, branch='main'):
        """上传文件到 GitHub"""
        if not GITHUB_AVAILABLE:
            return False, '未安装 PyGithub，请先 pip install pygithub'

        if not token:
            return False, '请输入 GitHub Token'

        try:
            g = Github(token)
            repo = g.get_repo(repo_name)
            filename = os.path.basename(file_path)

            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            try:
                contents = repo.get_contents(filename, ref=branch)
                repo.update_file(contents.path, f'Update {filename}', content, contents.sha, branch=branch)
            except Exception:
                repo.create_file(filename, f'Create {filename}', content, branch=branch)

            raw_url = f'https://raw.githubusercontent.com/{repo_name}/refs/heads/{branch}/{filename}'
            return True, raw_url
        except Exception as e:
            return False, f'上传失败: {str(e)}'

    def _copy_to_clipboard(self, text):
        """复制文本到剪贴板"""
        try:
            self.root.clipboard_clear()
            self.root.clipboard_append(text)
            self.root.update()
            return True
        except Exception as e:
            return False


def main():
    if DND_AVAILABLE:
        root = TkinterDnD.Tk()
    else:
        root = tk.Tk()

    # 设置主题
    try:
        style = ttk.Style()
        if 'vista' in style.theme_names():
            style.theme_use('vista')
        elif 'clam' in style.theme_names():
            style.theme_use('clam')
    except Exception:
        pass

    app = FileMergerApp(root)

    # 如果没有安装拖拽库，弹出提示
    if not DND_AVAILABLE:
        messagebox.showinfo(
            '提示',
            '未检测到 tkinterdnd2，拖拽功能不可用。\n\n'
            '可通过「添加文件」按钮选择文件。\n'
            '如需拖拽功能，请执行：\n'
            'pip install tkinterdnd2'
        )

    if not DOCX_AVAILABLE:
        print('[提示] 未安装 python-docx，Word 文档(.docx)将无法读取。安装命令：pip install python-docx')

    root.mainloop()


if __name__ == '__main__':
    main()
